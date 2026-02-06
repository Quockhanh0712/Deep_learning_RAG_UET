"""
Test DOCX file processing to debug upload issues
"""
import sys
from pathlib import Path

# Test if python-docx works
try:
    from docx import Document
    print("✓ python-docx imported successfully")
except ImportError as e:
    print(f"✗ Failed to import python-docx: {e}")
    sys.exit(1)

# Test creating a simple DOCX
try:
    doc = Document()
    doc.add_paragraph("Test paragraph 1")
    doc.add_paragraph("Test paragraph 2")
    
    test_path = Path(__file__).parent / "test_sample.docx"
    doc.save(str(test_path))
    print(f"✓ Created test DOCX: {test_path}")
    
    # Try reading it back
    doc2 = Document(str(test_path))
    text = "\n\n".join([para.text for para in doc2.paragraphs if para.text.strip()])
    print(f"✓ Read back test DOCX: {len(text)} chars")
    print(f"Content: {text}")
    
    # Clean up
    test_path.unlink()
    print("✓ Test DOCX processing successful!")
    
except Exception as e:
    print(f"✗ Error processing DOCX: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)

# Test SmartChunker with DOCX
print("\n" + "="*60)
print("Testing SmartRecursiveChunker with DOCX...")
print("="*60)

try:
    from smart_chunker import SmartRecursiveChunker
    
    # Create another test DOCX
    doc = Document()
    doc.add_paragraph("Điều 1. Phạm vi điều chỉnh")
    doc.add_paragraph("Luật này quy định về quyền và nghĩa vụ của công dân trong việc tham gia giao thông đường bộ.")
    doc.add_paragraph("Điều 2. Đối tượng áp dụng")
    doc.add_paragraph("Luật này áp dụng đối với mọi cá nhân, tổ chức tham gia giao thông trên lãnh thổ Việt Nam.")
    
    test_path = Path(__file__).parent / "test_sample.docx"
    doc.save(str(test_path))
    
    chunker = SmartRecursiveChunker(chunk_size=512)
    chunks = chunker.chunk_file(str(test_path))
    
    print(f"✓ SmartChunker processed DOCX: {len(chunks)} chunks")
    for i, chunk in enumerate(chunks):
        print(f"  Chunk {i+1}: {len(chunk.content)} chars - {chunk.content[:50]}...")
    
    # Clean up
    test_path.unlink()
    print("\n✓ All tests passed!")
    
except Exception as e:
    print(f"✗ SmartChunker failed: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)
