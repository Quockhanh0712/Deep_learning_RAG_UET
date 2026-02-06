"""
Smart Recursive Chunker for User Documents

Chiến lược xử lý file người dùng (PDF, DOCX, TXT):
1. Recursive Splitting: Paragraph → Sentence → Word
2. Overlap: 10-15% để đảm bảo tính liên tục
3. Context Injection: Chèn header [Nguồn: file.pdf | Trang X]

Author: Legal RAG System
Date: 2024-12
"""

import re
import os
import hashlib
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class Chunk:
    """Represents a text chunk with metadata"""
    content: str
    chunk_id: str
    source_file: str
    page_number: Optional[int] = None
    chunk_index: int = 0
    total_chunks: int = 0
    start_char: int = 0
    end_char: int = 0
    metadata: Dict = field(default_factory=dict)
    
    @property
    def header(self) -> str:
        """Generate context header for chunk"""
        if self.page_number:
            return f"[Nguồn: {self.source_file} | Trang {self.page_number}]"
        return f"[Nguồn: {self.source_file} | Đoạn {self.chunk_index + 1}/{self.total_chunks}]"
    
    @property
    def content_with_header(self) -> str:
        """Content with injected context header"""
        return f"{self.header}\n\n{self.content}"


class SmartRecursiveChunker:
    """
    Smart Recursive Text Chunker
    
    Features:
    - Recursive splitting by semantic boundaries
    - Configurable overlap for context preservation
    - Context header injection for source tracing
    - Support for PDF, DOCX, TXT files
    
    Usage:
        chunker = SmartRecursiveChunker(chunk_size=512, overlap=0.1)
        chunks = chunker.chunk_file("document.pdf")
    """
    
    # Separators ordered by priority (paragraph > sentence > phrase > word)
    SEPARATORS = [
        "\n\n\n",    # Triple newline (major sections)
        "\n\n",      # Double newline (paragraphs)
        "\n",        # Single newline
        ". ",        # Sentence end
        "? ",        # Question end
        "! ",        # Exclamation end
        "; ",        # Semicolon
        ", ",        # Comma
        " ",         # Space (last resort)
    ]
    
    def __init__(
        self,
        chunk_size: int = 512,
        overlap_ratio: float = 0.12,
        min_chunk_size: int = 100,
        max_chunk_size: int = 1024,
        inject_context: bool = True
    ):
        """
        Initialize chunker
        
        Args:
            chunk_size: Target chunk size in characters
            overlap_ratio: Overlap ratio (0.1 = 10%)
            min_chunk_size: Minimum chunk size
            max_chunk_size: Maximum chunk size
            inject_context: Whether to inject source context header
        """
        self.chunk_size = chunk_size
        self.overlap_ratio = overlap_ratio
        self.overlap_size = int(chunk_size * overlap_ratio)
        self.min_chunk_size = min_chunk_size
        self.max_chunk_size = max_chunk_size
        self.inject_context = inject_context
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters that might cause issues
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        # Normalize newlines
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        return text.strip()
    
    def _split_by_separator(
        self,
        text: str,
        separator: str
    ) -> List[str]:
        """Split text by separator, keeping separator with previous chunk"""
        if separator not in text:
            return [text]
        
        parts = text.split(separator)
        result = []
        
        for i, part in enumerate(parts):
            if i < len(parts) - 1:
                # Add separator back to the end of chunk (except last)
                result.append(part + separator)
            else:
                # Last part doesn't need separator
                if part.strip():
                    result.append(part)
        
        return [p for p in result if p.strip()]
    
    def _recursive_split(
        self,
        text: str,
        separator_index: int = 0
    ) -> List[str]:
        """
        Recursively split text into chunks
        
        Algorithm:
        1. If text fits in chunk_size, return as-is
        2. Try splitting by current separator
        3. If chunks still too big, recurse with next separator
        4. Combine small chunks to meet minimum size
        """
        # Base case: text fits
        if len(text) <= self.chunk_size:
            return [text] if text.strip() else []
        
        # Try current separator
        if separator_index >= len(self.SEPARATORS):
            # No more separators, force split by character
            return self._force_split(text)
        
        separator = self.SEPARATORS[separator_index]
        parts = self._split_by_separator(text, separator)
        
        # If no split happened, try next separator
        if len(parts) == 1:
            return self._recursive_split(text, separator_index + 1)
        
        # Process each part recursively if too large
        result = []
        for part in parts:
            if len(part) > self.chunk_size:
                # Part still too big, recurse with next separator
                result.extend(self._recursive_split(part, separator_index + 1))
            else:
                result.append(part)
        
        # Merge small adjacent chunks
        return self._merge_small_chunks(result)
    
    def _force_split(self, text: str) -> List[str]:
        """Force split long text by character count"""
        chunks = []
        start = 0
        
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            
            # Try to find a good break point
            if end < len(text):
                # Look for space near the end
                for i in range(end, max(start + self.min_chunk_size, end - 50), -1):
                    if text[i] == ' ':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end
        
        return chunks
    
    def _merge_small_chunks(self, chunks: List[str]) -> List[str]:
        """Merge adjacent small chunks"""
        if not chunks:
            return []
        
        result = []
        current = chunks[0]
        
        for chunk in chunks[1:]:
            combined = current + chunk
            
            if len(combined) <= self.chunk_size:
                current = combined
            else:
                if len(current) >= self.min_chunk_size:
                    result.append(current)
                    current = chunk
                else:
                    # Current is too small, force combine
                    current = combined
        
        if current.strip():
            result.append(current)
        
        return result
    
    def _add_overlap(self, chunks: List[str]) -> List[str]:
        """Add overlap between chunks for context continuity"""
        if len(chunks) <= 1 or self.overlap_size == 0:
            return chunks
        
        result = []
        
        for i, chunk in enumerate(chunks):
            if i == 0:
                result.append(chunk)
            else:
                # Get overlap from previous chunk
                prev_chunk = chunks[i - 1]
                overlap_text = prev_chunk[-self.overlap_size:] if len(prev_chunk) > self.overlap_size else prev_chunk
                
                # Find a good break point for overlap
                space_idx = overlap_text.find(' ')
                if space_idx > 0:
                    overlap_text = overlap_text[space_idx + 1:]
                
                # Add overlap prefix
                overlapped_chunk = f"...{overlap_text.strip()} {chunk}"
                result.append(overlapped_chunk)
        
        return result
    
    def chunk_text(
        self,
        text: str,
        source_file: str = "document",
        page_number: Optional[int] = None,
        add_overlap: bool = True
    ) -> List[Chunk]:
        """
        Chunk text into semantic pieces
        
        Args:
            text: Raw text to chunk
            source_file: Source filename for context injection
            page_number: Page number if from PDF
            add_overlap: Whether to add overlap between chunks
            
        Returns:
            List of Chunk objects with metadata
        """
        # Clean text
        clean_text = self._clean_text(text)
        
        if not clean_text:
            return []
        
        # Recursive split
        raw_chunks = self._recursive_split(clean_text)
        
        # Add overlap if requested
        if add_overlap:
            raw_chunks = self._add_overlap(raw_chunks)
        
        # Create Chunk objects with metadata
        chunks = []
        char_offset = 0
        
        for i, content in enumerate(raw_chunks):
            chunk_id = self._generate_chunk_id(source_file, i, content)
            
            chunk = Chunk(
                content=content,
                chunk_id=chunk_id,
                source_file=source_file,
                page_number=page_number,
                chunk_index=i,
                total_chunks=len(raw_chunks),
                start_char=char_offset,
                end_char=char_offset + len(content),
                metadata={
                    "source": source_file,
                    "page": page_number,
                    "chunk_index": i,
                    "total_chunks": len(raw_chunks)
                }
            )
            chunks.append(chunk)
            char_offset += len(content)
        
        return chunks
    
    def _generate_chunk_id(
        self,
        source: str,
        index: int,
        content: str
    ) -> str:
        """Generate unique chunk ID"""
        hash_input = f"{source}_{index}_{content[:50]}"
        return hashlib.md5(hash_input.encode()).hexdigest()[:16]
    
    def chunk_file(
        self,
        file_path: str,
        add_overlap: bool = True
    ) -> List[Chunk]:
        """
        Chunk a file (PDF, DOCX, TXT)
        
        Args:
            file_path: Path to file
            add_overlap: Whether to add overlap
            
        Returns:
            List of Chunk objects
        """
        path = Path(file_path)
        suffix = path.suffix.lower()
        source_name = path.name
        
        if suffix == '.pdf':
            return self._chunk_pdf(file_path, source_name, add_overlap)
        elif suffix == '.docx':
            return self._chunk_docx(file_path, source_name, add_overlap)
        elif suffix == '.doc':
            return self._chunk_doc_legacy(file_path, source_name, add_overlap)
        elif suffix == '.txt':
            return self._chunk_txt(file_path, source_name, add_overlap)
        else:
            logger.warning(f"Unsupported file type: {suffix}")
            return []
    
    def _chunk_pdf(
        self,
        file_path: str,
        source_name: str,
        add_overlap: bool
    ) -> List[Chunk]:
        """Extract and chunk PDF file"""
        try:
            import pymupdf  # PyMuPDF / fitz
        except ImportError:
            try:
                import fitz as pymupdf
            except ImportError:
                logger.error("PyMuPDF not installed. Run: pip install pymupdf")
                return []
        
        chunks = []
        
        try:
            doc = pymupdf.open(file_path)
            
            for page_num, page in enumerate(doc, 1):
                text = page.get_text()
                
                if text.strip():
                    page_chunks = self.chunk_text(
                        text=text,
                        source_file=source_name,
                        page_number=page_num,
                        add_overlap=add_overlap
                    )
                    chunks.extend(page_chunks)
            
            doc.close()
            logger.info(f"[CHUNKER] Extracted {len(chunks)} chunks from {source_name}")
            
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {e}")
        
        return chunks
    
    def _chunk_docx(
        self,
        file_path: str,
        source_name: str,
        add_overlap: bool
    ) -> List[Chunk]:
        """Extract and chunk DOCX file (Office Open XML format)"""
        try:
            from docx import Document
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return []
        
        try:
            logger.info(f"[CHUNKER] Opening DOCX file: {file_path}")
            doc = Document(file_path)
            
            # Extract all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            logger.info(f"[CHUNKER] Extracted {len(paragraphs)} paragraphs from DOCX")
            
            full_text = "\n\n".join(paragraphs)
            logger.info(f"[CHUNKER] Total text length: {len(full_text)} chars")
            
            if not full_text.strip():
                logger.warning(f"[CHUNKER] DOCX file has no text content: {file_path}")
                return []
            
            chunks = self.chunk_text(
                text=full_text,
                source_file=source_name,
                add_overlap=add_overlap
            )
            
            logger.info(f"[CHUNKER] Extracted {len(chunks)} chunks from {source_name}")
            return chunks
            
        except Exception as e:
            logger.error(f"[CHUNKER] Error reading DOCX {file_path}: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return []
    
    def _chunk_doc_legacy(
        self,
        file_path: str,
        source_name: str,
        add_overlap: bool
    ) -> List[Chunk]:
        """Extract and chunk legacy .DOC file (Word 97-2003 format)"""
        logger.info(f"[CHUNKER] Processing legacy .DOC file: {file_path}")
        
        # Try method 1: python-docx (might work for some .doc files that are actually .docx)
        try:
            from docx import Document
            logger.info("[CHUNKER] Attempting to open .doc as .docx...")
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            if paragraphs:
                full_text = "\n\n".join(paragraphs)
                logger.info(f"[CHUNKER] Successfully extracted {len(paragraphs)} paragraphs from .doc as .docx")
                
                chunks = self.chunk_text(
                    text=full_text,
                    source_file=source_name,
                    add_overlap=add_overlap
                )
                return chunks
        except Exception as e:
            logger.warning(f"[CHUNKER] Failed to open .doc as .docx: {e}")
        
        # Try method 2: docx2txt (lightweight alternative)
        try:
            import docx2txt
            logger.info("[CHUNKER] Attempting to extract text with docx2txt...")
            text = docx2txt.process(file_path)
            
            if text and text.strip():
                logger.info(f"[CHUNKER] Extracted {len(text)} chars with docx2txt")
                chunks = self.chunk_text(
                    text=text,
                    source_file=source_name,
                    add_overlap=add_overlap
                )
                return chunks
        except ImportError:
            logger.warning("[CHUNKER] docx2txt not installed")
        except Exception as e:
            logger.warning(f"[CHUNKER] docx2txt failed: {e}")
        
        # All methods failed - return error message as chunk for user feedback
        logger.error(f"[CHUNKER] Cannot process legacy .DOC file: {source_name}")
        logger.error("[CHUNKER] Solution: Please convert .doc to .docx format (open in Word and Save As .docx)")
        
        # Return empty list - will be handled by upload_file
        return []
    
    def _chunk_txt(
        self,
        file_path: str,
        source_name: str,
        add_overlap: bool
    ) -> List[Chunk]:
        """Extract and chunk TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            chunks = self.chunk_text(
                text=text,
                source_file=source_name,
                add_overlap=add_overlap
            )
            
            logger.info(f"[CHUNKER] Extracted {len(chunks)} chunks from {source_name}")
            return chunks
            
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {e}")
            return []


# ============================================================================
# Demo / Test
# ============================================================================

if __name__ == "__main__":
    # Demo text
    sample_text = """
    Điều 1. Phạm vi điều chỉnh
    
    Luật này quy định về quyền và nghĩa vụ của công dân trong việc tham gia giao thông đường bộ. 
    Các quy định cụ thể bao gồm việc sử dụng phương tiện giao thông, tuân thủ tín hiệu đèn, 
    và các biển báo giao thông trên đường.
    
    Điều 2. Đối tượng áp dụng
    
    Luật này áp dụng đối với mọi cá nhân, tổ chức tham gia giao thông trên lãnh thổ Việt Nam.
    Người nước ngoài cư trú tại Việt Nam cũng phải tuân thủ các quy định của Luật này.
    
    Điều 3. Giải thích từ ngữ
    
    Trong Luật này, các từ ngữ dưới đây được hiểu như sau:
    1. "Phương tiện giao thông" là các loại xe cơ giới và xe thô sơ được phép lưu hành.
    2. "Người tham gia giao thông" bao gồm người điều khiển phương tiện và người đi bộ.
    3. "Đường bộ" là phần mặt đường dành cho phương tiện và người đi bộ.
    """
    
    # Create chunker
    chunker = SmartRecursiveChunker(
        chunk_size=300,
        overlap_ratio=0.12,
        inject_context=True
    )
    
    # Chunk text
    chunks = chunker.chunk_text(
        text=sample_text,
        source_file="LuatGiaoThong.pdf",
        page_number=1
    )
    
    # Print results
    print("\n" + "=" * 60)
    print("SMART RECURSIVE CHUNKER - DEMO")
    print("=" * 60)
    
    for chunk in chunks:
        print(f"\n{'─' * 40}")
        print(f"📄 Chunk {chunk.chunk_index + 1}/{chunk.total_chunks}")
        print(f"   ID: {chunk.chunk_id}")
        print(f"   Length: {len(chunk.content)} chars")
        print(f"   Header: {chunk.header}")
        print(f"\n{chunk.content_with_header[:200]}...")
